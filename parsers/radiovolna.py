import requests
from bs4 import BeautifulSoup
import json
from datetime import datetime
import time
import os
import re
import yaml
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL  # Добавляем импорт для вертикального выравнивания
from docx.oxml.shared import OxmlElement, qn


class RadioVolnaParser:
    def __init__(self):
        self.base_url = "https://radiovolna.fm"
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        self.session = requests.Session()
        self.session.headers.update(self.headers)
        
        # Словарь для преобразования русских месяцев
        self.months_ru = {
            'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
            'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
            'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12
        }
        
        # Загружаем конфигурацию дат и поисковых слов
        self.start_date, self.end_date, self.search_list, self.reports_path, self.config = self.load_date_config()

    def load_date_config(self):
        """Загружает диапазон дат и список поисковых слов из config.yaml"""
        try:
            # Пробуем разные варианты пути к конфигу
            config_paths = [
                "parsers/config.yaml",
                "config.yaml", 
                os.path.join(os.path.dirname(__file__), "config.yaml")
            ]
            
            config_path = None
            for path in config_paths:
                if os.path.exists(path):
                    config_path = path
                    break
            
            if config_path:
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                    
                reports_section = config.get('reports', {})
                start_date = reports_section.get('start_date')
                end_date = reports_section.get('end_date')
                search_list = reports_section.get('search_list', [])
                reports_path = reports_section.get('path', 'data')
                
                if start_date and end_date:
                    print(f"📅 Диапазон дат из конфигурации: {start_date} — {end_date}")
                    print(f"🔍 Поисковые слова: {search_list}")
                    return start_date, end_date, search_list, reports_path, config
        
            # Значения по умолчанию, если конфиг не найден
            print("⚠️ Конфиг не найден, используются значения по умолчанию")
            return "2025-09-01", "2025-09-23", [], "data", {}
            
        except Exception as e:
            print(f"❌ Ошибка загрузки конфига: {e}")
            return "2025-09-01", "2025-09-23", [], "data", {}

    def check_keywords_in_title(self, title):
        """Проверяет наличие ключевых слов в заголовке с помощью регулярных выражений"""
        if not self.search_list or not title:
            return False, []
        
        found_keywords = []
        
        for keyword in self.search_list:
            # Создаем регулярное выражение для поиска слова (игнорируем регистр)
            pattern = re.compile(re.escape(keyword), re.IGNORECASE)
            if pattern.search(title):
                found_keywords.append(keyword)
        
        return len(found_keywords) > 0, found_keywords

    def is_date_in_range(self, date_str):
        """Проверяет, попадает ли дата в заданный диапазон"""
        if not date_str or date_str == "":
            return True  # Пропускаем новости без даты
        
        try:
            # Если дата уже в формате YYYY-MM-DD
            if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                return self.start_date <= date_str <= self.end_date
            return True
        except:
            return True  # В случае ошибки пропускаем

    def parse_russian_date(self, date_str):
        """Преобразует русскую дату в стандартный формат ISO"""
        if not date_str:
            return None
        
        try:
            # Ищем паттерн: число + месяц + год
            pattern = r'(\d{1,2})\s+(\w+)\s+(\d{4})'
            match = re.search(pattern, date_str)
            
            if match:
                day = int(match.group(1))
                month_name = match.group(2)
                year = int(match.group(3))
                
                # Преобразуем название месяца в номер
                if month_name in self.months_ru:
                    month = self.months_ru[month_name]
                    # Создаем объект datetime и возвращаем в формате ISO
                    date_obj = datetime(year, month, day)
                    return date_obj.strftime('%Y-%m-%d')
            
            return date_str  # Возвращаем оригинал, если не удалось распарсить
            
        except Exception as e:
            print(f"Ошибка парсинга даты '{date_str}': {e}")
            return date_str

    def get_page_content(self, url):
        """Получает содержимое страницы"""
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            return response.text
        except requests.RequestException as e:
            print(f"Ошибка загрузки {url}: {e}")
            return None

    def parse_new_materials_section(self, page_content):
        """Парсит только секцию 'Новые материалы'"""
        soup = BeautifulSoup(page_content, 'html.parser')
        news_items = []

        # Ищем секцию с заголовком "Новые материалы"
        sections = soup.find_all('section', class_='l-section')
        
        target_section = None
        for section in sections:
            title_elem = section.find('h2', class_='l-section__title')
            if title_elem and 'Новые материалы' in title_elem.get_text():
                target_section = section
                break
        
        if not target_section:
            print("Секция 'Новые материалы' не найдена")
            return news_items

        print("Найдена секция 'Новые материалы'")
        
        # Ищем все новостные блоки в этой секции
        news_blocks = target_section.find_all('div', class_='b-section-item')
        print(f"Найдено {len(news_blocks)} новостных блоков")
        
        for i, news_block in enumerate(news_blocks, 1):
            try:
                # Извлекаем заголовок
                title_elem = news_block.find('h3', class_='b-section-item__title')
                if not title_elem:
                    continue
                
                title_link = title_elem.find('a')
                if not title_link:
                    continue
                
                title = title_link.get_text(strip=True)
                href = title_link.get('href')
                
                # Формируем полный URL
                if href.startswith('/'):
                    full_url = self.base_url + href
                else:
                    full_url = href
                
                # Извлекаем дату
                date_text = ""
                meta_items = news_block.find_all('div', class_='b-meta-item')
                for meta_item in meta_items:
                    span = meta_item.find('span', class_='fa-clock-o')
                    if span:
                        date_span = span.find_next_sibling('span')
                        if date_span:
                            date_text = date_span.get_text(strip=True)
                        break
                
                # Извлекаем автора
                author = ""
                for meta_item in meta_items:
                    user_span = meta_item.find('span', class_='fa-user')
                    if user_span:
                        # Получаем текст после иконки пользователя
                        author_text = meta_item.get_text(strip=True)
                        author = author_text.strip()
                        break
                
                # Извлекаем URL изображения
                image_url = ""
                picture_div = news_block.find('div', class_='b-section-item__picture')
                if picture_div:
                    img = picture_div.find('img')
                    if img:
                        img_src = img.get('src') or img.get('data-src')
                        if img_src:
                            if img_src.startswith('/'):
                                image_url = self.base_url + img_src
                            else:
                                image_url = img_src
                
                # Преобразуем дату
                parsed_date = self.parse_russian_date(date_text)
                
                # Проверяем, попадает ли дата в диапазон
                if not self.is_date_in_range(parsed_date):
                    continue
                
                # Проверяем наличие ключевых слов в заголовке
                has_keywords, found_keywords = self.check_keywords_in_title(title)
                
                news_item = {
                    'title': title,
                    'url': full_url,
                    'date': parsed_date,
                    'date_original': date_text,  # Сохраняем оригинальную дату
                    'author': author,
                    'image_url': image_url,
                    'has_search_keywords': has_keywords,
                    'found_keywords': found_keywords,
                    'parsed_at': datetime.now().isoformat()
                }
                
                news_items.append(news_item)
                if has_keywords:
                    print(f"✓ Блок {i}: {title} 🔍 [{', '.join(found_keywords)}]")
                else:
                    print(f"✓ Блок {i}: {title}")
                
            except Exception as e:
                print(f"✗ Ошибка в блоке {i}: {e}")
                continue
        
        return news_items

    def save_to_json(self, data, filename):
        """Сохраняет данные в JSON файл в папку reports"""
        try:
            # Создаем папку reports, если её нет
            os.makedirs(self.reports_path, exist_ok=True)
            
            # Полный путь к файлу
            filepath = os.path.join(self.reports_path, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"Данные сохранены в {filepath}")
            return filepath
        except Exception as e:
            print(f"Ошибка сохранения: {e}")
            return None

    def create_news_table(self, doc, news_with_keywords, company_config):
        """Создает таблицу с новостями"""
        if not news_with_keywords:
            return None
        
        # Получаем домен и путь к лого из конфига
        site_domain = company_config.get('domain', 'radiovolna.fm')
        logo_path = company_config.get('logo_path', '../static/logos/logo_juka.png')
        
        # Создаем таблицу
        # Количество строк: 1 (заголовок) + (количество новостей * 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'  # Стиль таблицы с границами
        
        # Заголовки столбцов
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Наименование радиостанции/сайта'
        header_cells[1].text = 'Дата и время выхода в эфир'
        header_cells[2].text = 'Название программы'
        header_cells[3].text = 'Название информационного материала/ ссылка на публикацию'
        
        # Делаем заголовки жирными, центрируем первые 3 столбца и выравниваем по вертикали
        for i, cell in enumerate(header_cells):
            # Вертикальное центрирование для всех ячеек заголовка
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                # Центрируем первые 3 столбца горизонтально
                if i < 3:
                    paragraph.alignment = 1  # 1 = центрирование
        
        # Добавляем строки с новостями (по 2 строки на каждую новость)
        for news in news_with_keywords:
            # Форматируем дату
            formatted_date = ''
            if news.get('date'):
                try:
                    date_obj = datetime.strptime(news['date'], '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%d.%m.%y')
                except:
                    formatted_date = news.get('date', '')
            
            # ПЕРВАЯ СТРОКА: лого + дата + программа + заголовок
            row1_cells = table.add_row().cells
            
            # Устанавливаем вертикальное центрирование для всех ячеек строки
            for cell in row1_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Пытаемся вставить лого, если не получается - вставляем текст
            try:
                if os.path.exists(logo_path):
                    # Добавляем изображение лого
                    paragraph = row1_cells[0].paragraphs[0]
                    paragraph.alignment = 1  # Центрируем лого горизонтально
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(0.5))  # Размер лого
                else:
                    row1_cells[0].text = f"[ЛОГО]"
                    row1_cells[0].paragraphs[0].alignment = 1  # Центрируем горизонтально
            except Exception as e:
                print(f"⚠️ Не удалось вставить лого: {e}")
                row1_cells[0].text = f"[ЛОГО]"
                row1_cells[0].paragraphs[0].alignment = 1  # Центрируем горизонтально
            
            row1_cells[1].text = formatted_date
            row1_cells[1].paragraphs[0].alignment = 1  # Центрируем дату горизонтально
            
            row1_cells[2].text = '«Новости»'
            row1_cells[2].paragraphs[0].alignment = 1  # Центрируем программу горизонтально
            
            row1_cells[3].text = news['title']
            # Четвертый столбец остается с выравниванием по левому краю горизонтально
            
            # ВТОРАЯ СТРОКА: домен + дата + программа + ссылка
            row2_cells = table.add_row().cells
            
            # Устанавливаем вертикальное центрирование для всех ячеек строки
            for cell in row2_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            row2_cells[0].text = site_domain
            row2_cells[0].paragraphs[0].alignment = 1  # Центрируем домен горизонтально
            
            row2_cells[1].text = formatted_date
            row2_cells[1].paragraphs[0].alignment = 1  # Центрируем дату горизонтально
            
            row2_cells[2].text = '«Новости»'
            row2_cells[2].paragraphs[0].alignment = 1  # Центрируем программу горизонтально
            
            row2_cells[3].text = news.get('url', '')
            # Четвертый столбец остается с выравниванием по левому краю горизонтально
        
        print(f"📋 Создана таблица с {len(news_with_keywords)} новостями ({len(news_with_keywords) * 2} строк)")
        return table

    def replace_variables_in_paragraph(self, paragraph, company_config, news_with_keywords, doc):
        """Заменяет переменные в параграфе, включая $report"""
        period_text = self.format_period_ru(self.start_date, self.end_date)
        n_value = str((len(news_with_keywords) if news_with_keywords else 0) * 10)
        replacements = {
            '$pre_company_full': company_config.get('pre_company_full', ''),
            '$company_short': company_config.get('company_short', ''),
            '$post_company_full': company_config.get('post_company_full', ''),
            '$date': period_text,
            '$legal_address': company_config.get('legal_address', ''),
            '$inn_kpp': company_config.get('inn_kpp', ''),
            '$email': company_config.get('email', ''),
            '$ceo_name': company_config.get('ceo_name', ''),
            '$n': n_value,
        }
        
        # Получаем весь текст параграфа
        full_text = paragraph.text
        original_text = full_text
        
        # Проверяем, есть ли переменная $report
        if '$report' in full_text:
            print(f"🔍 Найдена переменная $report в тексте: {full_text}")
            
            # Заменяем остальные переменные в тексте
            text_parts = full_text.split('$report')
            text_before = text_parts[0] if len(text_parts) > 0 else ""
            text_after = text_parts[1] if len(text_parts) > 1 else ""
            
            # Заменяем переменные в частях до и после $report
            for var, value in replacements.items():
                text_before = text_before.replace(var, value)
                text_after = text_after.replace(var, value)
            
            # Очищаем текущий параграф и вставляем только текст до $report
            for run in paragraph.runs:
                run.text = ""
            
            if paragraph.runs:
                paragraph.runs[0].text = text_before
            else:
                paragraph.add_run(text_before)
            
            # Получаем позицию текущего параграфа в документе
            paragraph_element = paragraph._element
            parent = paragraph_element.getparent()
            paragraph_index = list(parent).index(paragraph_element)
            
            # Создаем таблицу с новостями
            if news_with_keywords:
                site_domain = company_config.get('domain', 'radiovolna.fm')
                logo_path = company_config.get('logo_path', '../static/logos/logo_juka.png')
                
                # Создаем таблицу
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Заголовки столбцов
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Наименование радиостанции/сайта'
                header_cells[1].text = 'Дата и время выхода в эфир'
                header_cells[2].text = 'Название программы'
                header_cells[3].text = 'Название информационного материала/ ссылка на публикацию'
                
                # Делаем заголовки жирными и центрируем
                for i, cell in enumerate(header_cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph_cell in cell.paragraphs:
                        for run in paragraph_cell.runs:
                            run.font.bold = True
                        if i < 3:
                            paragraph_cell.alignment = 1
                
                # Добавляем строки с новостями
                for news in news_with_keywords:
                    # Форматируем дату
                    formatted_date = ''
                    if news.get('date'):
                        try:
                            date_obj = datetime.strptime(news['date'], '%Y-%m-%d')
                            formatted_date = date_obj.strftime('%d.%m.%y')
                        except:
                            formatted_date = news.get('date', '')
                    
                    # ПЕРВАЯ СТРОКА: лого + дата + программа + заголовок
                    row1_cells = table.add_row().cells
                    for cell in row1_cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Лого
                    try:
                        if os.path.exists(logo_path):
                            paragraph_logo = row1_cells[0].paragraphs[0]
                            paragraph_logo.alignment = 1
                            run = paragraph_logo.runs[0] if paragraph_logo.runs else paragraph_logo.add_run()
                            run.add_picture(logo_path, width=Inches(0.5))
                        else:
                            row1_cells[0].text = "[ЛОГО]"
                            row1_cells[0].paragraphs[0].alignment = 1
                    except Exception as e:
                        row1_cells[0].text = "[ЛОГО]"
                        row1_cells[0].paragraphs[0].alignment = 1
                    
                    row1_cells[1].text = formatted_date
                    row1_cells[1].paragraphs[0].alignment = 1
                    row1_cells[2].text = '«Новости»'
                    row1_cells[2].paragraphs[0].alignment = 1
                    row1_cells[3].text = news['title']
                    
                    # ВТОРАЯ СТРОКА: домен + дата + программа + ссылка
                    row2_cells = table.add_row().cells
                    for cell in row2_cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    row2_cells[0].text = site_domain
                    row2_cells[0].paragraphs[0].alignment = 1
                    row2_cells[1].text = formatted_date
                    row2_cells[1].paragraphs[0].alignment = 1
                    row2_cells[2].text = '«Новости»'
                    row2_cells[2].paragraphs[0].alignment = 1
                    row2_cells[3].text = news.get('url', '')
                
                # Вставляем таблицу СРАЗУ ПОСЛЕ текущего параграфа
                table_element = table._element
                parent.insert(paragraph_index + 1, table_element)
                
                print(f"📋 Таблица вставлена на место $report с {len(news_with_keywords)} новостями")
            
            # Если есть текст после $report, добавляем его в новый параграф
            if text_after.strip():
                new_paragraph = doc.add_paragraph(text_after)
                new_paragraph_element = new_paragraph._element
                parent.insert(paragraph_index + 2, new_paragraph_element)
            
            return True
        
        # Обычная замена переменных (если нет $report)
        has_variables = any(var in full_text for var in replacements.keys())
        
        if has_variables:
            print(f"🔍 Найден текст с переменными: {full_text}")
            
            new_text = full_text
            for var, value in replacements.items():
                if var in new_text:
                    new_text = new_text.replace(var, value)
                    print(f"🔄 Заменено: {var} → {value}")
            
            if new_text != original_text:
                for run in paragraph.runs:
                    run.text = ""
                
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                
                print(f"✅ Обновлен параграф: {new_text}")
        
        return False

    def replace_variables_in_cell(self, cell, company_config, news_with_keywords, doc):
        """Заменяет переменные в ячейке таблицы"""
        for paragraph in cell.paragraphs:
            self.replace_variables_in_paragraph(paragraph, company_config, news_with_keywords, doc)

    def process_template_document(self, news_with_keywords, company_name='Южка'):
        """Обрабатывает template.docx: копирует шаблон, заменяет переменные включая $report"""
        try:
            # Получаем конфигурацию компании
            company_config = self.config.get('reports', {}).get('docs', {}).get(company_name, {})
            if not company_config:
                print(f"❌ Конфигурация для '{company_name}' не найдена")
                return None
            
            print(f"🔧 Конфигурация '{company_name}':")
            for key, value in company_config.items():
                print(f"   {key}: {value}")
            
            # Пути к .docx файлам
            template_paths = [
                "static/template.docx",
                "../static/template.docx"
            ]
            
            template_path = None
            for path in template_paths:
                if os.path.exists(path):
                    template_path = path
                    break
            
            if not template_path:
                print("❌ Файл template.docx не найден")
                return None
            
            print(f"📄 Используем шаблон: {template_path}")
            
            # Создаем папку reports, если её нет
            os.makedirs(self.reports_path, exist_ok=True)
            
            # Генерируем имя выходного файла
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"processed_document_{company_name}_{timestamp}.docx"
            output_path = os.path.join(self.reports_path, output_filename)
            
            # Копируем .docx файл целиком
            print("📋 Копируем template.docx...")
            shutil.copy2(template_path, output_path)
            doc = Document(output_path)
            
            print("🔄 Ищем и заменяем переменные...")
            
            # Заменяем переменные во всех параграфах
            # Важно: проходим по копии списка, так как мы можем изменять документ
            paragraphs = list(doc.paragraphs)
            for paragraph in paragraphs:
                if paragraph.text.strip():
                    self.replace_variables_in_paragraph(paragraph, company_config, news_with_keywords, doc)
            
            # Заменяем переменные в таблицах
            for table_num, table in enumerate(doc.tables):
                print(f"�� Обрабатываем таблицу {table_num + 1}")
                for row_num, row in enumerate(table.rows):
                    for cell_num, cell in enumerate(row.cells):
                        if cell.text.strip():
                            self.replace_variables_in_cell(cell, company_config, news_with_keywords, doc)
            
            print("✅ Обработка переменных завершена")
            
            # Сохраняем обработанный документ
            doc.save(output_path)
            print(f"�� Документ сохранен: {output_path}")
            
            return output_path
            
        except Exception as e:
            print(f"❌ Ошибка обработки документа: {e}")
            import traceback
            traceback.print_exc()
            return None

    def print_results(self, news_items):
        """Выводит результаты парсинга"""
        for i, item in enumerate(news_items, 1):
            print(f"{i}. {item['title']}")
            if item.get('page_number'):
                print(f"   Страница: {item['page_number']}")
            if item['author']:
                print(f"   Автор: {item['author']}")
            if item['date']:
                print(f"   Дата: {item['date']}")
            print(f"   URL: {item['url']}")
            print()

    def format_period_ru(self, start_iso: str, end_iso: str) -> str:
        """Возвращает период в формате: 15 по 21 августа 2025 (с учетом месяцев и года).
        Ожидает ISO строки YYYY-MM-DD."""
        try:
            start_dt = datetime.strptime(start_iso, "%Y-%m-%d")
            end_dt = datetime.strptime(end_iso, "%Y-%m-%d")
        except Exception:
            # В случае ошибки просто вернем исходные строки
            return f"{start_iso} — {end_iso}"
        months_gen = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]
        if start_dt.year == end_dt.year:
            if start_dt.month == end_dt.month:
                # 15 по 21 августа 2025
                month_name = months_gen[end_dt.month - 1]
                return f"{start_dt.day} по {end_dt.day} {month_name} {end_dt.year}"
            else:
                # 25 июля по 03 августа 2025
                start_month = months_gen[start_dt.month - 1]
                end_month = months_gen[end_dt.month - 1]
                return f"{start_dt.day} {start_month} по {end_dt.day} {end_month} {end_dt.year}"
        else:
            # 28 декабря 2024 по 03 января 2025
            start_month = months_gen[start_dt.month - 1]
            end_month = months_gen[end_dt.month - 1]
            return f"{start_dt.day} {start_month} {start_dt.year} по {end_dt.day} {end_month} {end_dt.year}"


def main():
    """Основная функция"""
    parser = RadioVolnaParser()
    
    # Базовый URL
    base_url = "https://radiovolna.fm/"
    
    print("Запуск парсинга секции 'Новые материалы' с автоматической остановкой по датам...")
    print(f"Базовый URL: {base_url}")
    print(f"Диапазон дат: {parser.start_date} — {parser.end_date}")
    print("=" * 60)
    
    all_news = []
    seen_titles = set()  # Для отслеживания уже найденных заголовков
    page_num = 1  # Начинаем с первой страницы
    
    # Цикл с автоматической остановкой по датам
    while True:
        target_url = f"{base_url}?PAGEN_4={page_num}"
        
        print(f"\n📄 СТРАНИЦА {page_num}")
        print(f"URL: {target_url}")
        print("-" * 40)
        
        # Получаем содержимое страницы
        page_content = parser.get_page_content(target_url)
        if not page_content:
            print(f"❌ Не удалось загрузить страницу {page_num}")
            page_num += 1
            continue
        
        # Парсим секцию "Новые материалы"
        news_data = parser.parse_new_materials_section(page_content)
        
        if news_data:
            # Фильтруем дубли по заголовку и проверяем даты
            new_items = []
            items_before_start = []  # Новости старше start_date
            
            for item in news_data:
                if item['title'] not in seen_titles:
                    item['page_number'] = page_num
                    
                    # Проверяем дату
                    if item['date'] and item['date'] < parser.start_date:
                        items_before_start.append(item)
                    else:
                        new_items.append(item)
                        seen_titles.add(item['title'])
            
            print(f"✅ Найдено {len(news_data)} новостей на странице {page_num}")
            if len(new_items) != len(news_data):
                filtered_count = len(news_data) - len(new_items)
                print(f"   🔄 Исключено {filtered_count} (дубли + вне диапазона)")
            print(f"   ➕ Добавлено {len(new_items)} уникальных новостей в диапазоне")
            
            # Добавляем новости в диапазоне
            all_news.extend(new_items)
            
            # Выводим только уникальные заголовки с этой страницы
            if new_items:
                print("\nНовые заголовки в диапазоне:")
                for i, item in enumerate(new_items, 1):
                    if item.get('has_search_keywords', False):
                        keywords_str = ', '.join(item.get('found_keywords', []))
                        print(f"  {i}. {item['title']} ({item['date']}) 🔍 [{keywords_str}]")
                    else:
                        print(f"  {i}. {item['title']} ({item['date']})")
            
            # Проверяем, есть ли новости старше start_date
            if items_before_start:
                print(f"🏁 Найдены новости старше {parser.start_date}, завершаем парсинг")
                break
                
            # Если на странице нет новостей в диапазоне, но есть новости - продолжаем
            if not new_items and news_data:
                print("   ⏭️ Все новости на странице вне диапазона, проверяем следующую")
        else:
            print(f"❌ Новости не найдены на странице {page_num}")
            # Если несколько страниц подряд пустые - останавливаемся
            if page_num > 25:  # После 25 страницы останавливаемся если пусто
                print("🏁 Слишком много пустых страниц, завершаем парсинг")
                break
        
        # Инкрементируем номер страницы
        page_num += 1
        
        # Пауза между запросами
        time.sleep(1)
        
        # Защита от бесконечного цикла
        if page_num > 100:
            print("⚠️ Достигнут лимит страниц (100), завершаем парсинг")
            break
    
    # Общие результаты
    print(f"\n{'='*60}")
    print(f"ОБЩИЕ РЕЗУЛЬТАТЫ")
    print(f"{'='*60}")
    
    if all_news:
        print(f"Всего обработано страниц: {page_num - 1}")
        print(f"Всего найдено новостей в диапазоне: {len(all_news)}")
        
        # Показываем диапазон дат в результатах
        dates = [item['date'] for item in all_news if item['date']]
        if dates:
            min_date = min(dates)
            max_date = max(dates)
            print(f"Диапазон дат в результатах: {min_date} — {max_date}")
        
        # Статистика по ключевым словам
        news_with_keywords = [item for item in all_news if item.get('has_search_keywords', False)]
        print(f"Новостей с ключевыми словами: {len(news_with_keywords)} из {len(all_news)}")
        
        if news_with_keywords:
            # Подсчитываем частоту каждого ключевого слова
            keyword_count = {}
            for item in news_with_keywords:
                for keyword in item.get('found_keywords', []):
                    keyword_count[keyword] = keyword_count.get(keyword, 0) + 1
            
            print("Статистика ключевых слов:")
            for keyword, count in sorted(keyword_count.items(), key=lambda x: x[1], reverse=True):
                print(f"  '{keyword}': {count} упоминаний")
        
        # Обрабатываем template.docx
        print(f"\n{'='*60}")
        print("ОБРАБОТКА ДОКУМЕНТА")
        print(f"{'='*60}")
        
        doc_file = parser.process_template_document(news_with_keywords)
        
        if doc_file:
            print(f"✅ Документ успешно создан: {doc_file}")
        else:
            print("❌ Ошибка создания документа")
        
    else:
        print("❌ Новости не найдены ни на одной странице")


if __name__ == "__main__":
    main()
