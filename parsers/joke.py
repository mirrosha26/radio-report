import requests
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import time
import os
import re
import yaml
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL


class JokeParser:
    def __init__(self):
        self.base_url = "https://www.astrobl.ru"
        self.list_url = f"{self.base_url}/press/news"
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0 Safari/537.36'
        }
        self.months_ru = {
            'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
            'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
            'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12
        }
        self.start_date, self.end_date, self.search_list, self.reports_path, self.config = self.load_date_config()

    def load_date_config(self):
        try:
            config_paths = [
                "parsers/config.yaml",
                "config.yaml",
                os.path.join(os.path.dirname(__file__), "config.yaml")
            ]
            config_path = None
            for path in config_paths:
                if os.path.exists(path):
                    config_path = path
                    break
            if config_path:
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                reports = config.get('reports', {})
                return (
                    reports.get('start_date'),
                    reports.get('end_date'),
                    reports.get('search_list', []),
                    reports.get('path', '../reports/'),
                    config,
                )
            return "2025-09-01", "2025-09-23", [], "../reports/", {}
        except Exception:
            return "2025-09-01", "2025-09-23", [], "../reports/", {}

    def format_period_ru(self, start_iso: str, end_iso: str) -> str:
        try:
            start_dt = datetime.strptime(start_iso, "%Y-%m-%d")
            end_dt = datetime.strptime(end_iso, "%Y-%m-%d")
        except Exception:
            return f"{start_iso} — {end_iso}"
        months_gen = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]
        if start_dt.year == end_dt.year:
            if start_dt.month == end_dt.month:
                month_name = months_gen[end_dt.month - 1]
                return f"{start_dt.day} по {end_dt.day} {month_name} {end_dt.year}"
            else:
                return f"{start_dt.day} {months_gen[start_dt.month-1]} по {end_dt.day} {months_gen[end_dt.month-1]} {end_dt.year}"
        else:
            return f"{start_dt.day} {months_gen[start_dt.month-1]} {start_dt.year} по {end_dt.day} {months_gen[end_dt.month-1]} {end_dt.year}"

    def check_keywords_in_title(self, title: str):
        if not self.search_list or not title:
            return False, []
        found = []
        for kw in self.search_list:
            # Строгое совпадение с учетом регистра и границ слова/фразы
            pattern = r"(?<!\\w)" + re.escape(kw) + r"(?!\\w)"
            if re.search(pattern, title):  # без IGNORECASE
                found.append(kw)
        return len(found) > 0, found

    def fetch_article_text(self, url: str) -> str:
        html = self.get(url)
        if not html:
            return ''
        try:
            soup = BeautifulSoup(html, 'html.parser')
            article = soup.select_one('.news-text, [itemprop="articleBody"]')
            if not article:
                article = soup.select_one('.content__body_main')
            return article.get_text(' ', strip=True) if article else ''
        except Exception:
            return ''

    def check_keywords_in_text(self, text: str):
        if not self.search_list or not text:
            return False, []
        found = []
        for kw in self.search_list:
            pattern = r"(?<!\\w)" + re.escape(kw) + r"(?!\\w)"
            if re.search(pattern, text):  # без IGNORECASE
                found.append(kw)
        return len(found) > 0, found

    def get(self, url: str):
        try:
            resp = requests.get(url, headers=self.headers, timeout=15)
            resp.raise_for_status()
            return resp.text
        except Exception as e:
            print(f"Ошибка загрузки {url}: {e}")
            return None

    def build_period_param(self) -> str:
        # DD.MM.YYYY+-+DD.MM.YYYY
        try:
            sd = datetime.strptime(self.start_date, "%Y-%m-%d").strftime("%d.%m.%Y")
            ed = datetime.strptime(self.end_date, "%Y-%m-%d").strftime("%d.%m.%Y")
        except Exception:
            ed = datetime.now().strftime("%d.%m.%Y")
            sd = (datetime.now() - timedelta(days=3)).strftime("%d.%m.%Y")
        return f"{sd}+-+{ed}"

    def parse_date_text(self, text: str) -> str:
        # Примеры: "Сегодня в 13:13", "Вчера в 9:28", "Вторник, 9:21", "23 сентября 2025, 14:55"
        t = text.strip()
        today = datetime.now()
        if t.lower().startswith("сегодня"):
            return today.strftime("%Y-%m-%d")
        if t.lower().startswith("вчера"):
            return (today - timedelta(days=1)).strftime("%Y-%m-%d")
        # День недели, например: "Вторник, 9:21" → ближайший прошедший вторник
        weekdays = {
            'понедельник': 0,
            'вторник': 1,
            'среда': 2,
            'четверг': 3,
            'пятница': 4,
            'суббота': 5,
            'воскресенье': 6,
        }
        mw = re.match(r"^([А-Яа-я]+)\s*(?:,|\s)\s*\d{1,2}:\d{2}", t)
        if mw:
            wd_name = mw.group(1).lower()
            if wd_name in weekdays:
                target = weekdays[wd_name]
                diff = (today.weekday() - target) % 7
                if diff == 0:
                    diff = 7  # всегда прошлая неделя, если совпал с сегодняшним днём недели
                dt = today - timedelta(days=diff)
                return dt.strftime("%Y-%m-%d")
        m = re.search(r"(\d{1,2})\s+(\w+)\s+(\d{4})", t)
        if m:
            day = int(m.group(1))
            month_name = m.group(2).lower()
            year = int(m.group(3))
            month = self.months_ru.get(month_name)
            if month:
                try:
                    return datetime(year, month, day).strftime("%Y-%m-%d")
                except Exception:
                    return t
        return t

    def parse_list_page(self, html: str):
        soup = BeautifulSoup(html, 'html.parser')
        articles = soup.find_all('article', class_='news')
        items = []
        for art in articles:
            try:
                a = art.find('a')
                if not a:
                    continue
                title_el = art.select_one('.title h5')
                title = title_el.get_text(strip=True) if title_el else a.get('title') or ''
                href = a.get('href') or ''
                url = href if href.startswith('http') else (self.base_url + href)
                img_div = art.select_one('.img')
                image_url = ''
                if img_div and img_div.has_attr('style'):
                    style = img_div['style']
                    m = re.search(r"url\(\"?([^\)\"]+)\"?\)", style)
                    if m:
                        src = m.group(1)
                        image_url = src if src.startswith('http') else (self.base_url + src)
                date_span = art.select_one('.date span[title="Дата публикации"]') or art.select_one('.date span')
                date_text = date_span.get_text(strip=True) if date_span else ''
                parsed_date = self.parse_date_text(date_text)
                has_kw, found_kw = self.check_keywords_in_title(title)
                items.append({
                    'title': title,
                    'url': url,
                    'date': parsed_date,
                    'date_original': date_text,
                    'author': '',
                    'image_url': image_url,
                    'has_search_keywords': has_kw,
                    'found_keywords': found_kw,
                    'has_kw_title': has_kw,
                    'found_keywords_title': list(found_kw) if found_kw else [],
                    'parsed_at': datetime.now().isoformat(),
                })
            except Exception as e:
                print(f"Ошибка статьи: {e}")
                continue
        return items

    def create_news_table(self, doc, news_with_keywords, company_config):
        if not news_with_keywords:
            return None
        site_domain = company_config.get('domain', 'astrobl.ru')
        logo_path = company_config.get('logo_path', '../static/logos/logo_juka.png')
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        headers = [
            'Наименование радиостанции',
            'Дата и время выхода в эфир',
            'Название программы',
            'Название информационного материала',
        ]
        for i, txt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = txt
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                if i < 3:
                    p.alignment = 1
        def _format_display_date(date_str: str, original_str: str) -> str:
            if date_str:
                # ISO yyyy-mm-dd
                m = re.match(r"^(\d{4}-\d{2}-\d{2})$", date_str)
                if m:
                    try:
                        return datetime.strptime(m.group(1), '%Y-%m-%d').strftime('%d.%m.%y')
                    except Exception:
                        pass
                # ISO with time yyyy-mm-dd HH:MM:SS
                m2 = re.match(r"^(\d{4}-\d{2}-\d{2})\s+\d{2}:\d{2}:\d{2}$", date_str)
                if m2:
                    try:
                        return datetime.strptime(m2.group(1), '%Y-%m-%d').strftime('%d.%m.%y')
                    except Exception:
                        pass
            # Попытаться распарсить оригинал
            if original_str:
                try:
                    reparsed = self.parse_date_text(original_str)
                    mm = re.match(r"^(\d{4}-\d{2}-\d{2})$", reparsed or '')
                    if mm:
                        return datetime.strptime(mm.group(1), '%Y-%m-%d').strftime('%d.%m.%y')
                except Exception:
                    pass
            return ''
        for news in news_with_keywords:
            formatted_date = _format_display_date(news.get('date') or '', news.get('date_original') or '')
            r1 = table.add_row().cells
            for c in r1:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            try:
                if os.path.exists(logo_path):
                    p = r1[0].paragraphs[0]
                    p.alignment = 1
                    run = p.runs[0] if p.runs else p.add_run()
                    run.add_picture(logo_path, width=Inches(0.5))
                else:
                    r1[0].text = '[ЛОГО]'
                    r1[0].paragraphs[0].alignment = 1
            except Exception:
                r1[0].text = '[ЛОГО]'
                r1[0].paragraphs[0].alignment = 1
            r1[1].text = formatted_date
            r1[1].paragraphs[0].alignment = 1
            r1[2].text = '«Новости»'
            r1[2].paragraphs[0].alignment = 1
            r1[3].text = news['title']
            
        return table

    def replace_variables_in_paragraph(self, paragraph, company_config, news_with_keywords, doc):
        period_text = self.format_period_ru(self.start_date, self.end_date)
        n_value = str((len(news_with_keywords) if news_with_keywords else 0) * 10)
        replacements = {
            '$pre_company_full': company_config.get('pre_company_full', ''),
            '$company_short': company_config.get('company_short', ''),
            '$post_company_full': company_config.get('post_company_full', ''),
            '$date': period_text,
            '$legal_address': company_config.get('legal_address', ''),
            '$inn_kpp': company_config.get('inn_kpp', ''),
            '$email': company_config.get('email', ''),
            '$ceo_name': company_config.get('ceo_name', ''),
            '$ogrn': company_config.get('ogrn', ''),
            '$n': n_value,
        }
        text = paragraph.text
        if '$report' in text:
            parts = text.split('$report')
            before = parts[0]
            after = parts[1] if len(parts) > 1 else ''
            for k, v in replacements.items():
                before = before.replace(k, v)
                after = after.replace(k, v)
            for r in paragraph.runs:
                r.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = before
            else:
                paragraph.add_run(before)
            pe = paragraph._element
            parent = pe.getparent()
            idx = list(parent).index(pe)
            table = self.create_news_table(doc, news_with_keywords, company_config or {})
            if table is not None:
                parent.insert(idx + 1, table._element)
            if after.strip():
                new_p = doc.add_paragraph(after)
                parent.insert(idx + 2, new_p._element)
            return True
        new_text = text
        changed = False
        for k, v in replacements.items():
            if k in new_text:
                new_text = new_text.replace(k, v)
                changed = True
        if changed:
            for r in paragraph.runs:
                r.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
        return changed

    def replace_variables_in_cell(self, cell, company_config, news_with_keywords, doc):
        for p in cell.paragraphs:
            self.replace_variables_in_paragraph(p, company_config, news_with_keywords, doc)

    def process_template_document(self, news_with_keywords, company_name='Юмор'):
        try:
            company_config = self.config.get('reports', {}).get('docs', {}).get(company_name, {})
            template_paths = [
                "static/template2.docx",
                "../static/template2.docx",
            ]
            template_path = None
            for path in template_paths:
                if os.path.exists(path):
                    template_path = path
                    break
            if not template_path:
                print("❌ template2.docx не найден")
                return None
            os.makedirs(self.reports_path, exist_ok=True)
            out_name = f"parsed_{company_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            out_path = os.path.join(self.reports_path, out_name)
            shutil.copy2(template_path, out_path)
            doc = Document(out_path)
            paragraphs = list(doc.paragraphs)
            for p in paragraphs:
                if p.text.strip():
                    self.replace_variables_in_paragraph(p, company_config, news_with_keywords, doc)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            self.replace_variables_in_cell(cell, company_config, news_with_keywords, doc)
            doc.save(out_path)
            print(f"✅ Документ сохранен: {out_path}")
            return out_path
        except Exception as e:
            print(f"❌ Ошибка документа: {e}")
            return None

    def run(self):
        print("Запуск парсинга astrobl.ru...")
        print(f"Период: {self.start_date} — {self.end_date}")
        period_param = self.build_period_param()
        try:
            start_dt = datetime.strptime(self.start_date, "%Y-%m-%d")
        except Exception:
            start_dt = None
        try:
            end_dt = datetime.strptime(self.end_date, "%Y-%m-%d")
        except Exception:
            end_dt = None
        page = 0
        all_items = []
        seen = set()
        while True:
            url = f"{self.list_url}?page={page}&period={period_param}"
            print(f"Страница {page}: {url}")
            html = self.get(url)
            if not html:
                print("❌ Нет HTML, прерываем")
                break
            items = self.parse_list_page(html)
            if not items:
                print("🏁 Контент закончился, выходим")
                break
            # Фильтрация по периоду на уровне страницы и ранний выход
            page_has_newer = False
            page_has_in_range = False
            page_has_older = False
            filtered_items = []
            for it in items:
                item_dt = None
                d = it.get('date') or ''
                if re.match(r"\d{4}-\d{2}-\d{2}", d):
                    try:
                        item_dt = datetime.strptime(d, '%Y-%m-%d')
                    except Exception:
                        item_dt = None
                if item_dt is not None and end_dt is not None and item_dt > end_dt:
                    page_has_newer = True
                    continue
                if item_dt is not None and start_dt is not None and item_dt < start_dt:
                    page_has_older = True
                    continue
                page_has_in_range = True
                filtered_items.append(it)
            if not page_has_in_range and page_has_older:
                print("🏁 Достигли начала периода, выходим")
                break
            print(f"Заголовки на странице (в диапазоне): {len(filtered_items)}")
            for it in filtered_items:
                print(f"  - {it.get('date','')} | {it.get('title','')}")
            new_cnt = 0
            for it in filtered_items:
                key = (it['title'], it['url'])
                if key not in seen:
                    seen.add(key)
                    all_items.append(it)
                    new_cnt += 1
            print(f"Добавлено уникальных: {new_cnt}")
            page += 1
            time.sleep(0.8)
            if page > 100:
                print("⚠️ Лимит страниц 100")
                break
        print(f"Итого собрано: {len(all_items)}")
        # Обогащаем совпадения за счет текста статьи, только если в заголовке не найдено
        for it in all_items:
            try:
                if not it.get('has_search_keywords'):
                    print(f"Проверяю текст статьи: {it.get('url','')}")
                    body_text = self.fetch_article_text(it['url'])
                    
                    has_kw_body, found_kw_body = self.check_keywords_in_text(body_text)
                    if has_kw_body:
                        print(f"  ✔ Найдено в тексте: {', '.join(found_kw_body)}")
                        it['has_search_keywords'] = True
                        existed = set(it.get('found_keywords') or [])
                        for kw in found_kw_body:
                            if kw not in existed:
                                existed.add(kw)
                        it['found_keywords'] = list(existed)
                    else:
                        print("  ✖ Ключевых слов в тексте не найдено")
            except Exception as e:
                print(f"  ⚠ Ошибка при разборе текста: {e}")
        news_with_kw = [x for x in all_items if x.get('has_search_keywords')]
        print(f"В отчет попадут ({len(news_with_kw)}):")
        for it in news_with_kw:
            print(f"  → {it.get('date','')} | {it.get('title','')} | keys: {', '.join(it.get('found_keywords') or [])}")
        # Сортируем от самых старых к самым новым по ISO-дате (YYYY-MM-DD)
        def _date_key(item):
            d = item.get('date') or ''
            if re.match(r"\d{4}-\d{2}-\d{2}", d):
                try:
                    return datetime.strptime(d, '%Y-%m-%d')
                except Exception:
                    return datetime.min
            return datetime.min
        news_with_kw.sort(key=_date_key)
        print(f"Совпадений по ключевым словам: {len(news_with_kw)}")
        self.process_template_document(news_with_kw, company_name='Юмор')


def main():
    JokeParser().run()


if __name__ == "__main__":
    main()